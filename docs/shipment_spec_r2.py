from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def sf(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    sf(run, bold=True, size=14, color=(0x1A,0x1D,0x23))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), '3B5BDB')
    pBdr.append(bot); pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    sf(run, bold=True, size=11, color=(0x3B,0x5B,0xDB))

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    sf(run, size=10)

def bul(text, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent * 0.6)
    run = p.add_run(text)
    sf(run, size=10)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run('※ ' + text)
    sf(run, size=9, color=(0x5A,0x60,0x70))

def tbl(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'),'EEF2FF'); tcPr.append(shd)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); sf(run, bold=True, size=9, color=(0x3B,0x5B,0xDB))
    for ri, rd in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(rd):
            c = row.cells[ci]; c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = c.paragraphs[0].add_run(str(val)); sf(run, size=9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ══════════════════════ 표지 ══════════════════════
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run('입출하관리 화면'), bold=True, size=22, color=(0x1A,0x1D,0x23))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run('기능 명세서'), bold=True, size=16, color=(0x3B,0x5B,0xDB))
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run('ERP 개발 요청용  |  2026-05  |  R2'), size=10, color=(0x5A,0x60,0x70))
doc.add_page_break()


# ══════════════════════ 1. 개요 (R1 그대로 유지) ══════════════════════
h1('1. 개요')
body('생산과제정보가 등록된 과제의 각 공정 (FAB / BUMP / CP / ASSEMBLY / DPS / FT) 단계별 투입/출하 이력을 관리하는 화면입니다.')
body('과제(프로젝트) 단위로 관리되며, 각 공정 단계에서 IN/OUT 기록을 등록·조회할 수 있습니다.')
doc.add_paragraph()
tbl(
    ['항목','내용'],
    [
        ['화면명','입출하관리'],
        ['대상 공정','FAB / BUMP / CP / ASSEMBLY / FT / DPS'],
        ['관리 단위','MPW(EA), Wafer·BUMP·CP(PCS & EA), Assembly·FT·DPS(EA)'],
        ['주요 기능','IN/OUT 등록, 잔여 추적, YLD 자동계산, Excel 업로드, 메모장 파일 파싱'],
    ],
    col_widths=[4,12]
)


# ══════════════════════ 2. 화면 구성 (R1 그대로 유지) ══════════════════════
h1('2. 화면 구성')
body('화면은 좌측 과제 목록(사이드바)과 우측 상세 패널로 구성됩니다.')
doc.add_paragraph()
tbl(
    ['영역','설명'],
    [
        ['과제 목록 (좌측)','등록된 과제를 검색·선택. 과제명·AL Code 검색 가능'],
        ['과제 정보 바 (상단)','선택된 과제의 AL Code·TM Code·Net Die·고객사 등 기본 정보 표시'],
        ['공정 탭','해당 과제의 SCM Flow(예: FAB → CP)를 탭으로 표시. 탭 클릭 시 해당 단계 기록 조회 (해당 step은 생산과제정보 등록 페이지에서 등록됨)'],
        ['기록 테이블 (우측 메인)','선택 공정의 IN/OUT 기록 목록. IN 추가, OUT 등록, 편집, 삭제 가능'],
    ],
    col_widths=[4,12]
)


# ══════════════════════ 3. 과제 분류 (R2 수정) ══════════════════════
h1('3. 과제 분류 및 공정별 관리 단위')

body('과제의 SCM Flow는 고정되어 있지 않으며, 생산과제정보 등록 화면에서 과제별로 자유롭게 조합합니다.')
body('(예: FAB 단독 / FAB→CP / FAB→ASSEMBLY→FT / FAB→BUMP→CP→ASSEMBLY→FT→DPS 등)')
body('입출하관리 화면의 공정 탭은 해당 과제에 등록된 SCM Flow 기준으로 자동 생성됩니다.')
doc.add_paragraph()

body('공정별 관리 단위는 아래 기준으로 구분됩니다.')
doc.add_paragraph()

tbl(
    ['공정','분류','관리 단위','비고'],
    [
        ['FAB','MPW','EA','IN=OUT 수량 고정. Invoice(TSMC) 파싱. 잔여/YLD 개념 없음. 이후 ASSEMBLY 연결 가능'],
        ['FAB','Wafer','PCS + EA','Lot당 최대 25장. Wafer No. 자동 할당(#1~N). 잔여 장수 실시간 추적'],
        ['BUMP','Wafer','PCS + EA','Wafer No. 직접 입력. Wafer별 Good Die 수량 입력 그리드'],
        ['CP','Wafer','PCS + EA','Wafer No. 직접 입력. Wafer별 Good Die 수량 입력 그리드'],
        ['ASSEMBLY','—','EA','이전 단계 OUT EA 기준 직접 입력. Reject·YLD 자동계산'],
        ['FT','—','EA','EA 직접 입력. Reject·YLD 자동계산'],
        ['DPS','—','EA','EA 직접 입력. Reject·YLD 자동계산'],
    ],
    col_widths=[2.5, 2.5, 3, 8]
)
note('PCS = Wafer 수량, EA = Die 수량 (PCS × Net Die). Net Die는 생산과제정보에서 과제별 등록.')


# ══════════════════════ 4. FAB 단계 기능 ══════════════════════
h1('4. FAB 단계 기능')

h2('4-1. FAB IN 등록')
tbl(
    ['구분','입력 항목'],
    [
        ['MPW','Ver / 날짜 / Start Qty(EA) / Remark'],
        ['Wafer','Ver / 날짜 / Lot ID / Wafer 수량(최대 25) → Wafer No. 자동 할당(#1~N) / Remark'],
    ],
    col_widths=[3,13]
)
note('Lot ID 형식: 영숫자 6자리 + "." + 2자리 (예: ABC123.01)')

h2('4-2. FAB OUT 등록')
body('[MPW]')
bul('TSMC Invoice 파싱 지원 — .txt 파일 업로드 또는 텍스트 붙여넣기')
bul('파싱 결과: Invoice No. / 품목별 Chip ID·수량 자동 추출')
bul('상태·처리방향 선택 후 저장. FAB 이후 ASSEMBLY 공정이 있는 경우 처리방향을 "다음step"으로 지정')
doc.add_paragraph()
body('[Wafer]')
bul('한 FAB IN에 대해 복수의 OUT 분배 등록 가능 (예: #1~3 고객출하 / #4~6 보관 / #7 scrap)')
bul('Wafer Pool 시각화 — 이미 처리된 Wafer는 비활성 표시')
bul('분배별 입력: Wafer No.(범위 표기 지원) / 처리방향 / 상태 / Reject / YLD% / Remark')
bul('IN 행에 잔여 장수 배지 실시간 표시 (잔여 N장 / 완출하 / 미출하)')
note('Wafer No. 범위 입력: 쉼표·~ 사용 (예: 1~3,5,7~10). Excel 날짜 오인식 방지를 위해 ~ 권장')

h2('4-3. 처리방향 옵션')
body('다음step / 고객출하 / 보관 / scrap / 기타')


# ══════════════════════ 5. BUMP / CP 단계 기능 ══════════════════════
h1('5. BUMP / CP 단계 기능')
body('Wafer 단위로 관리됩니다. PCS(Wafer 수량)와 EA(Good Die 수량)를 함께 기록합니다.')
body('Lot 단위로 등록하며, 각 Wafer별 Good Die 수량을 개별 입력합니다.')

h2('5-1. Lot IN 등록')
tbl(
    ['입력 항목','설명'],
    [
        ['Lot ID','직접 입력. 이전 단계 Lot ID 자동 가이드 제공 (클릭 시 자동 입력)'],
        ['Wafer No.','범위 입력 (예: 1~3,5~6). FAB scrap 등으로 인한 비연속 번호 지원'],
        ['In(PCS)','Wafer 수량 — Wafer No. 입력 기준 자동계산'],
        ['In(EA)','Net Die × In(PCS) 자동계산. Net Die 미등록 시 직접 입력'],
        ['입고일 / 시작일 / 종료일','날짜 입력'],
        ['상태 / 처리방향 / Remark',''],
    ],
    col_widths=[4,12]
)

h2('5-2. Lot OUT — Wafer별 Good Die 입력')
bul('Wafer No. 입력 시 Wafer별 입력 그리드 자동 생성')
bul('각 Wafer마다 Out Good Die(EA) 수량 및 Scrap 여부 입력')
bul('Out(PCS) = Scrap 제외 Wafer 수, Out(EA) = Good Die 합산 — 자동계산')
bul('Reject = In(EA) − Out(EA),  YLD% = Out(EA) ÷ In(EA) × 100 — 자동계산')
bul('Scrap 체크 시 해당 Wafer 행 비활성화')


# ══════════════════════ 6. ASSEMBLY / FT / DPS 단계 기능 ══════════════════════
h1('6. ASSEMBLY / FT / DPS 단계 기능')
body('EA 단위로만 관리합니다. Wafer 번호 개념 없음.')
body('이전 단계 OUT EA를 기준으로 IN 수량을 직접 입력하며, OUT 입력 시 Reject·YLD가 자동계산됩니다.')

tbl(
    ['입력 항목','설명'],
    [
        ['Lot ID','직접 입력. 이전 단계 Lot ID 자동 가이드 제공'],
        ['In(EA)','직접 입력 (이전 단계 OUT EA 기준)'],
        ['Out(EA)','직접 입력'],
        ['Reject','In(EA) − Out(EA) 자동계산'],
        ['YLD%','Out(EA) ÷ In(EA) × 100 자동계산'],
        ['입고일 / 시작일 / 종료일','날짜 입력'],
        ['상태 / 처리방향 / Remark',''],
    ],
    col_widths=[4,12]
)
note('MPW FAB 이후 ASSEMBLY가 연결되는 경우, FAB OUT의 총 EA가 ASSEMBLY IN(EA) 기준이 됩니다.')


# ══════════════════════ 7. Excel 업로드 기능 ══════════════════════
h1('7. Excel 업로드 기능')
body('공정 단계별 Excel 일괄 등록을 지원합니다. .xlsx / .xls / .csv 파일 업로드 가능.')
doc.add_paragraph()
tbl(
    ['모드','헤더 컬럼','대상 공정'],
    [
        ['FAB IN',        'Ver | 날짜 | Lot ID | Wafer Qty (1-25) | Remark',                                         'FAB (Wafer)'],
        ['FAB OUT',       'Ver | 날짜 | Lot ID | Wafer No. | 처리방향 | 상태 | Remark',                               'FAB (Wafer)'],
        ['Lot IN',        'Lot ID | Wafer Qty | 입고일 | 시작일 | 종료일 | In(PCS) | In(EA) | 상태 | 처리방향 | Remark','BUMP / CP'],
        ['Lot OUT',       'Lot ID | Wafer No. | Out Good Die (EA) | Is Scrap (Y/N) | Remark',                         'BUMP / CP'],
        ['ASSY/FT/DPS IN', 'Lot ID | In(EA) | 입고일 | 시작일 | 종료일 | 상태 | 처리방향 | Remark',                  'ASSEMBLY / FT / DPS'],
        ['ASSY/FT/DPS OUT','Lot ID | Out(EA) | 상태 | 처리방향 | Remark',                                            'ASSEMBLY / FT / DPS'],
    ],
    col_widths=[3, 9.5, 3.5]
)
bul('업로드 전 템플릿 미리보기 및 헤더 클립보드 복사 기능 제공')
bul('파일 선택 후 데이터 미리보기 확인 → 저장 버튼으로 일괄 등록')
bul('날짜 컬럼: Excel 날짜 시리얼 자동 변환 (YYYY-MM-DD)')
bul('Wafer No. 범위: ~ 기호 사용 권장 (예: 1~3,5~20) — - 기호는 Excel이 날짜로 인식할 수 있음')


# ══════════════════════ 8. 공통 기능 ══════════════════════
h1('8. 공통 기능')
tbl(
    ['기능','설명'],
    [
        ['검색','과제명·AL Code 키워드 검색'],
        ['편집','등록된 IN/OUT/Lot 레코드 수정 (기존 값 자동 로드)'],
        ['삭제','레코드 삭제 (확인 팝업 후 처리)'],
        ['Wafer별 상세 보기','BUMP·CP Lot 행에서 ▼ 버튼 클릭 시 Wafer별 Good Die·Scrap 현황 인라인 표시'],
        ['잔여 배지','FAB IN(Wafer) 행에 잔여 N장 / 완출하 / 미출하 상태 실시간 표시'],
        ['Lot ID 가이드','Lot 등록 시 이전 단계 Lot ID 자동 안내 — 클릭 시 자동 입력'],
    ],
    col_widths=[4,12]
)


# ══════════════════════ 9. 용어 정의 ══════════════════════
h1('9. 용어 정의')
tbl(
    ['용어','정의'],
    [
        ['Lot ID',   '공정 단위 식별자. 형식: 영숫자 6자리 + "." + 2자리 (예: ABC123.01)'],
        ['Net Die',  '웨이퍼 1장당 Good Die 수량 (과제 기준 고정값, 생산과제정보에서 등록)'],
        ['In(PCS)',  '투입 Wafer 수량'],
        ['In(EA)',   '투입 Die 수량 (= PCS × Net Die). ASSEMBLY·FT·DPS는 직접 입력'],
        ['Out(PCS)', '출하 Wafer 수량 (Scrap 제외)'],
        ['Out(EA)',  '출하 Good Die 수량. ASSEMBLY·FT·DPS는 직접 입력'],
        ['Reject',   'In(EA) − Out(EA)'],
        ['YLD%',     'Out(EA) ÷ In(EA) × 100'],
        ['처리방향', '다음step / 고객출하 / 보관 / scrap / 기타'],
        ['SCM Flow', '과제별 공정 순서. 생산과제정보 등록 화면에서 과제마다 자유롭게 구성'],
    ],
    col_widths=[3.5,12.5]
)


out = '/home/user/mj-automation-dashboard/docs/입출하관리_기능명세서_R2.docx'
doc.save(out)
print('saved:', out)
