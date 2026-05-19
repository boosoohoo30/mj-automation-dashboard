from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
sec = doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7)
sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
doc.styles['Normal'].font.name='맑은 고딕'
doc.styles['Normal'].font.size=Pt(10)

def sf(run, bold=False, size=10, color=None):
    run.font.name='맑은 고딕'; run.font.size=Pt(size); run.font.bold=bold
    if color: run.font.color.rgb=RGBColor(*color)

def h1(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)
    run=p.add_run(text); sf(run,bold=True,size=14,color=(0x1A,0x1D,0x23))
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'4'); bot.set(qn('w:color'),'3B5BDB'); pBdr.append(bot); pPr.append(pBdr)

def h2(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    run=p.add_run(text); sf(run,bold=True,size=11,color=(0x3B,0x5B,0xDB))

def h3(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
    run=p.add_run(text); sf(run,bold=True,size=10,color=(0x1A,0x1D,0x23))

def body(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(text); sf(run,size=10)

def bul(text, indent=1):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
    p.paragraph_format.left_indent=Cm(indent*0.6)
    run=p.add_run(text); sf(run,size=10)

def note(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(4)
    p.paragraph_format.left_indent=Cm(0.4); run=p.add_run('※ '+text); sf(run,size=9,color=(0x5A,0x60,0x70))

def tbl(headers, rows, col_widths=None):
    t=doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        tcPr=c._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'),'EEF2FF'); tcPr.append(shd)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(h); sf(run,bold=True,size=9,color=(0x3B,0x5B,0xDB))
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1]
        for ci,val in enumerate(rd):
            c=row.cells[ci]; c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            run=c.paragraphs[0].add_run(str(val)); sf(run,size=9)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph(); return t

def shade_row(table, row_idx, fill_hex):
    row=table.rows[row_idx]
    for cell in row.cells:
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'),fill_hex); tcPr.append(shd)

# ══════════════════ 표지 ══════════════════
doc.add_paragraph(); doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run('입출하관리 화면'),bold=True,size=22,color=(0x1A,0x1D,0x23))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run('기능 명세서'),bold=True,size=16,color=(0x3B,0x5B,0xDB))
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
sf(p.add_run('ERP 개발 요청용  |  2026-05  |  R5'),size=10,color=(0x5A,0x60,0x70))
doc.add_page_break()

# ══════════════════ 1. 개요 ══════════════════
h1('1. 개요')
body('생산과제정보가 등록된 과제의 각 공정 (FAB / BUMP / CP / ASSEMBLY / DPS / FT) 단계별 투입/출하 이력을 관리하는 화면입니다.')
body('과제(프로젝트) 단위로 관리되며, 각 공정 단계에서 IN/OUT 기록을 등록·조회할 수 있습니다.')
doc.add_paragraph()
tbl(['항목','내용'],[
    ['화면명','입출하관리'],
    ['대상 공정','FAB / BUMP / CP / ASSEMBLY / FT / DPS'],
    ['관리 단위','MPW(EA), Wafer·BUMP·CP(PCS & EA), Assembly·FT·DPS(EA)'],
    ['주요 기능','IN/OUT 등록, 잔여 추적, YLD 자동계산, Excel 업로드, MPW/Wafer Invoice 파싱(txt)'],
],col_widths=[4,12])

# ══════════════════ 2. 화면 구성 ══════════════════
h1('2. 화면 구성')
body('화면은 좌측 과제 목록(사이드바)과 우측 상세 패널로 구성됩니다.')
doc.add_paragraph()
tbl(['영역','설명'],[
    ['과제 목록 (좌측)','등록된 과제를 검색·선택. 과제명·AL Code 검색 가능'],
    ['과제 정보 바 (상단)','선택된 과제의 AL Code·TM Code·Net Die·고객사 등 기본 정보 표시'],
    ['공정 탭','해당 과제의 SCM Flow를 탭으로 표시. 탭 클릭 시 해당 단계 기록 조회 (step은 생산과제정보 등록 페이지에서 등록됨)'],
    ['기록 테이블 (우측 메인)','선택 공정의 IN/OUT 기록 목록. IN 추가, OUT 등록, 편집, 삭제 가능'],
],col_widths=[4,12])

# ══════════════════ 3. 과제 분류 ══════════════════
h1('3. 과제 분류 및 공정별 관리 단위')
body('과제의 SCM Flow는 고정되어 있지 않으며, 생산과제정보 등록 화면에서 과제별로 자유롭게 조합합니다.')
body('(예: FAB 단독 / FAB→CP / FAB→ASSEMBLY→FT / FAB→BUMP→CP→ASSEMBLY→FT→DPS 등)')
body('입출하관리 화면의 공정 탭은 해당 과제에 등록된 SCM Flow 기준으로 자동 생성됩니다.')
doc.add_paragraph()
body('공정은 FAB과 Post-FAB으로 구분되며, Post-FAB은 관리 단위에 따라 두 가지로 나뉩니다.')
doc.add_paragraph()
t=tbl(['공정','구분','관리 단위','설명'],[
    ['FAB (MPW)','FAB','EA','IN = OUT 수량 고정. Invoice 파싱(txt). 잔여/YLD 없음. 이후 ASSEMBLY 연결 가능'],
    ['FAB (Wafer)','FAB','PCS + EA','Lot당 최대 25장. Wafer No. 자동 할당. 잔여 장수 실시간 추적'],
    ['BUMP','Post-FAB','PCS + EA','Wafer No. 직접 입력. Wafer별 Good Die 입력 그리드'],
    ['CP','Post-FAB','PCS + EA','Wafer No. 직접 입력. Wafer별 Good Die 입력 그리드'],
    ['ASSEMBLY','Post-FAB','EA','EA 직접 입력. Reject·YLD 자동계산'],
    ['FT','Post-FAB','EA','EA 직접 입력. Reject·YLD 자동계산'],
    ['DPS','Post-FAB','EA','EA 직접 입력. Reject·YLD 자동계산'],
],col_widths=[2.5,2.5,2.5,8.5])
shade_row(t,1,'FFF8F0'); shade_row(t,2,'FFF8F0')
shade_row(t,3,'F0F4FF'); shade_row(t,4,'F0F4FF')
shade_row(t,5,'F3F9ED'); shade_row(t,6,'F3F9ED'); shade_row(t,7,'F3F9ED')
note('PCS = Wafer 수량, EA = Die 수량 (PCS × Net Die). Net Die는 생산과제정보에서 과제별 등록.')

# ══════════════════ 4. FAB 단계 기능 ══════════════════
h1('4. FAB 단계 기능')

h2('4-1. FAB IN 등록')
tbl(['구분','입력 항목'],[
    ['MPW','Ver / 날짜 / Start Qty(EA) / Remark'],
    ['Wafer','Ver / 날짜 / Lot 목록(다중) / Remark'],
],col_widths=[3,13])
bul('Wafer FAB IN: Lot을 여러 개 입력 가능 — 동일 Invoice에 Lot이 복수이면 한 번에 등록')
bul('각 Lot 행: Lot ID + 수량(Lot당 최대 25장) → Wafer No. #1~N Lot별 독립 자동 할당')
bul('+ Lot 추가 버튼으로 동적 행 추가. 총 수량 = 각 Lot 합산 (예: 3 Lot × 25장 = 75장)')
note('Lot당 최대 25장 (TSMC 1 Lot 기준). 예: 75장 Invoice → 3 Lot × 25장으로 1건 등록')

h2('4-2. FAB OUT 등록')
h3('[MPW]')
bul('TSMC Invoice 파싱 지원 — .txt 파일 업로드 또는 텍스트 직접 붙여넣기 후 파싱 실행')
bul('파싱 결과 테이블: Invoice No. / Invoice Date / AWB No. / PO No. / 과제명 / TM Code / 수량(EA) / Lot ID / Wafer No. 자동 추출')
bul('파싱 합계 + Carton 합계 각각 별도 표시 (일치 여부 지시자 없음)')
bul('상태 / 처리방향 선택 후 저장. FAB 이후 ASSEMBLY 연결 시 처리방향을 "다음step"으로 지정')
bul('Invoice 파싱 상세 로직은 7절 참조')
doc.add_paragraph()
body('상태 옵션: 진행중 / 완료 / Scrap')

doc.add_paragraph()
h3('[Wafer]')
bul('TSMC Invoice 파싱 지원 (선택사항) — .txt 파일 업로드 또는 텍스트 붙여넣기 후 파싱 실행')
bul('파싱 완료 후 "분배에 적용" 버튼: Wafer No. 포함 시 분배 블록에 자동 채움, 미포함 시 수기 입력')
bul('한 FAB IN에 대해 복수의 OUT 분배 등록 가능 (예: #1~3 고객출하 / #4~6 보관 / #7 scrap)')
bul('분배별 입력: Wafer No.(범위 표기 지원) / 상태 / 처리방향 / Remark')
bul('IN 행에 잔여 장수 배지 실시간 표시 (잔여 N장 / 출하 / 미출하)')
bul('[다중 Lot] FAB IN이 다중 Lot이면 Pool을 Lot별 섹션으로 분리 표시 (각 Lot 독립 #1-N)')
bul('[다중 Lot] 저장 시 각 분배의 TSMC Lot ID가 OUT 레코드 Lot ID로 자동 설정')
bul('[검증] 단일 Lot: 분배 총 수량 ≤ FAB IN 수량. 다중 Lot: Lot별 분배 수량 ≤ 해당 Lot 수량')
note('Wafer No. 범위 입력: 쉼표·~ 사용 (예: 1~3,5,7~10). Excel 날짜 오인식 방지를 위해 ~ 권장')
doc.add_paragraph()
body('상태 옵션: 진행중 / 완료 / Scrap')

h2('4-3. 처리방향 옵션')
body('다음step / 고객출하 / 보관 / scrap / 기타')

# ══════════════════ 5. Post-FAB 단계 기능 ══════════════════
h1('5. Post-FAB 단계 기능')
body('FAB 이후 공정(BUMP / CP / ASSEMBLY / FT / DPS)을 총칭합니다.')
body('관리 단위에 따라 ① PCS+EA 방식 (BUMP·CP) 과 ② EA 방식 (ASSEMBLY·FT·DPS) 으로 구분됩니다.')

h2('5-1. BUMP / CP — Wafer 단위 (PCS + EA)')
body('Wafer 번호를 직접 입력하며, 각 Wafer별 Good Die 수량을 개별 입력합니다.')
doc.add_paragraph()

h3('Lot IN 입력 항목')
tbl(['항목','설명'],[
    ['Lot ID','직접 입력. 이전 단계 Lot ID 자동 가이드 제공 (클릭 시 자동 입력)'],
    ['Wafer No.','범위 입력 (예: 1~3,5~6). FAB scrap 등으로 인한 비연속 번호 지원'],
    ['In(PCS)','Wafer 수량 — Wafer No. 입력 기준 자동계산'],
    ['In(EA)','Net Die × In(PCS) 자동계산. Net Die 미등록 시 직접 입력'],
    ['입고일 / 시작일 / 종료일','날짜 입력'],
    ['상태 / 처리방향 / Remark',''],
],col_widths=[4,12])

h3('Lot OUT — Wafer별 Good Die 입력')
bul('Wafer No. 입력 시 Wafer별 입력 그리드 자동 생성')
bul('각 Wafer마다 Out Good Die(EA) 수량 및 Scrap 여부 입력')
bul('Out(PCS) = Scrap 제외 Wafer 수 / Out(EA) = Good Die 합산 — 자동계산')
bul('Reject = In(EA) − Out(EA),  YLD% = Out(EA) ÷ In(EA) × 100 — 자동계산')
bul('Scrap 체크 시 해당 Wafer 행 비활성화')

doc.add_paragraph()
h2('5-2. ASSEMBLY / FT / DPS — EA 단위')
body('Wafer 번호 개념 없음. 이전 단계 OUT EA를 기준으로 IN 수량을 직접 입력합니다.')
doc.add_paragraph()
tbl(['입력 항목','설명'],[
    ['Lot ID','직접 입력. 이전 단계 Lot ID 자동 가이드 제공'],
    ['In(EA)','직접 입력 (이전 단계 OUT EA 기준)'],
    ['Out(EA)','직접 입력'],
    ['Reject','In(EA) − Out(EA) 자동계산'],
    ['YLD%','Out(EA) ÷ In(EA) × 100 자동계산'],
    ['입고일 / 시작일 / 종료일','날짜 입력'],
    ['상태 / 처리방향 / Remark',''],
],col_widths=[4,12])
note('MPW FAB 이후 ASSEMBLY가 연결되는 경우, FAB OUT의 총 EA가 ASSEMBLY IN(EA) 기준이 됩니다.')

# ══════════════════ 6. Excel 업로드 ══════════════════
h1('6. Excel 업로드 기능')
body('공정 단계별 Excel 일괄 등록을 지원합니다. .xlsx / .xls / .csv 파일 업로드 가능.')
doc.add_paragraph()
tbl(['모드','헤더 컬럼','대상 공정'],[
    ['FAB IN',  'Ver | 날짜 | Lot ID | Wafer Qty (1-25) | Remark',                                           'FAB (Wafer)'],
    ['FAB OUT', 'Ver | 날짜 | Lot ID | Wafer No. | 처리방향 | 상태 | Remark',                                 'FAB (Wafer)'],
    ['Lot IN',  'Lot ID | Wafer Qty | 입고일 | 시작일 | 종료일 | In(PCS) | In(EA) | 상태 | 처리방향 | Remark', 'BUMP / CP'],
    ['Lot OUT', 'Lot ID | Wafer No. | Out Good Die (EA) | Is Scrap (Y/N) | Remark',                           'BUMP / CP'],
    ['EA IN',   'Lot ID | In(EA) | 입고일 | 시작일 | 종료일 | 상태 | 처리방향 | Remark',                     'ASSEMBLY / FT / DPS'],
    ['EA OUT',  'Lot ID | Out(EA) | 상태 | 처리방향 | Remark',                                               'ASSEMBLY / FT / DPS'],
],col_widths=[2.5,10,3.5])
bul('업로드 전 템플릿 미리보기 및 헤더 클립보드 복사 기능 제공')
bul('파일 선택 후 데이터 미리보기 확인 → 저장 버튼으로 일괄 등록')
bul('날짜 컬럼: Excel 날짜 시리얼 자동 변환 (YYYY-MM-DD)')
bul('Wafer No. 범위: ~ 기호 사용 권장 (예: 1~3,5~20) — - 기호는 Excel이 날짜로 인식할 수 있음')

# ══════════════════ 7. TSMC Invoice 파싱 ══════════════════
h1('7. TSMC Invoice 파싱 기능 (메모장 .txt)')
body('MPW 및 Wafer FAB OUT 등록 시 TSMC 발행 Invoice(.txt)를 파싱하여 출하 정보를 자동 추출합니다.')
body('.txt 파일 업로드 또는 텍스트 붙여넣기 → 파싱 실행 → 결과 미리보기 확인 → 저장 순서로 동작합니다.')

h2('7-1. 공통 헤더 추출')
tbl(['추출 항목','Invoice 내 위치','비고'],[
    ['Invoice No.','INVOICE NO.: 뒤 값','저장 시 기록'],
    ['Invoice Date','INVOICE DATE: 뒤 값','참고용 표시'],
    ['AWB No.','HAWB NO. : 뒤 값','참고용 표시'],
],col_widths=[4,6,6])

h2('7-2. MPW — Chip 블록 파싱 (1블록 = 1행)')
body('AL- 로 시작하는 항목 번호 행을 기준으로 블록을 탐지하고, 이후 하위 라인에서 아래 값을 추출합니다.')
doc.add_paragraph()
tbl(['추출 항목','Invoice 내 위치'],[
    ['PO No.','CUSTOMER ORDER 컬럼 첫 번째 값 (예: AL-MPW-20251125-11)'],
    ['과제명 (name)','PO No. 다음 줄 왼쪽 값'],
    ['TM Code','과제명과 같은 줄 오른쪽 값'],
    ['Qty','단위(EA) 앞 숫자'],
    ['단위','EA 고정'],
    ['Lot ID','EA 다음 줄, # 앞 값 (예: G6X997.1#1 → Lot ID = G6X997.1)'],
    ['Wafer No.','# 뒤 숫자 (예: G6X997.1#1 → Wafer No. = 1), 표시 시 #1 형식'],
],col_widths=[4,12])
note('v10 구현 완료: 두 자리 아이템 번호(10+) 및 페이지 브레이크 대응. Lot ID / Wafer No. 분리 처리.')

h2('7-3. Wafer — Lot ID 행 파싱 (1행 = 1 Lot)')
body('Wafer Invoice는 MPW와 다른 포맷입니다. * 기호로 Lot ID와 수량을 구분합니다.')
doc.add_paragraph()
tbl(['추출 항목','Invoice 내 위치'],[
    ['규격','12" WAFER, 8" WAFER 등 — 규격 문자열은 파싱에 사용하지 않음 (AL- 패턴 기준으로만 아이템 탐지)'],
    ['PO No.','CUSTOMER ORDER 컬럼 첫 번째 값 (예: AL-RTO-20240115-02)'],
    ['과제명 / TM Code','PO No. 다음 줄 좌·우 값'],
    ['Lot ID','* 앞 값 (예: P6BM75.00)'],
    ['Qty','* 뒤 숫자'],
    ['단위','PCS 고정'],
    ['Wafer No.','Lot ID 뒤 #범위 값이 있으면 파싱 (예: P6BM75#1-6, T60F45#2,5,6), 없으면 빈칸 (수기 입력)'],
],col_widths=[4,12])
note('v10 구현 완료: 8"/12" Wafer 모두 대응. 1개 아이템 내 복수 Lot(예: 6개 Lot), 복수 아이템 모두 파싱.')
note('Wafer No. 범위 포맷: 연속(#1-6 → [1,2,3,4,5,6]), 비연속(#2,5,6 → [2,5,6]), 혼합(#1,3-5 → [1,3,4,5]) 지원.')

h2('7-4. 합계 표시')
body('파싱 완료 후 파싱 합산값과 GRAND TOTAL(Invoice 내 기재 값)을 각각 별도로 표시합니다.')
body('일치 여부 지시자는 표시하지 않습니다.')
doc.add_paragraph()
tbl(['표시 항목','설명'],[
    ['파싱 합계','파싱된 항목 수량 합산값 (자동 계산)'],
    ['Carton 합계','Invoice 내 GRAND TOTAL 또는 Carton QTY 합산값 (자동 추출)'],
],col_widths=[4,12])

h2('7-5. 구현 상태 (v10 R5 기준)')
t=tbl(['항목','상태','비고'],[
    ['Invoice No. 추출','✅','저장 시 기록'],
    ['Invoice Date / AWB No. 추출','✅','참고용 표시'],
    ['MPW Chip 블록 탐지 (AL- 패턴)','✅','두 자리 번호·페이지 브레이크 대응'],
    ['과제명 / TM Code / PO No. 추출','✅','—'],
    ['Qty / 단위 추출','✅','MPW: EA, Wafer: PCS'],
    ['Lot ID + Wafer No. 분리','✅','# 기준 분리 처리'],
    ['Wafer 타입 Invoice 파싱','✅','* 기준 Lot ID / Qty 추출, Wafer No. 범위 파싱'],
    ['Wafer No. 범위 자동 분배 적용','✅','# 있으면 자동 채움, 없으면 수기 입력'],
    ['GRAND TOTAL 표시','✅','파싱 합계 + Carton 합계 각각 별도 표시'],
    ['FAB IN 다중 Lot 입력','✅','+ Lot 추가, Lot별 최대 25장, 총합 저장'],
    ['다중 Lot Pool 시각화','✅','FAB OUT 모달에서 Lot별 섹션 (#1-N 독립)'],
    ['다중 Lot 수량 검증','✅','Lot별 OUT 합계 ≤ FAB IN 해당 Lot 수량'],
    ['OUT lotId 자동 설정','✅','각 분배의 tmLotId → OUT 레코드 Lot ID'],
],col_widths=[5.5,1.5,9])
for i in range(1,14): shade_row(t,i,'F3F9ED')

# ══════════════════ 8. 공통 기능 ══════════════════
h1('8. 공통 기능')
tbl(['기능','설명'],[
    ['검색','과제명·AL Code 키워드 검색'],
    ['편집','등록된 IN/OUT/Lot 레코드 수정 (기존 값 자동 로드)'],
    ['삭제','레코드 삭제 (확인 팝업 후 처리)'],
    ['Wafer별 상세 보기','BUMP·CP Lot 행에서 ▼ 버튼 클릭 시 Wafer별 Good Die·Scrap 현황 인라인 표시'],
    ['잔여 배지','FAB IN(Wafer) 행에 잔여 N장 / 출하 / 미출하 상태 실시간 표시'],
    ['Lot ID 가이드','Lot 등록 시 이전 단계 Lot ID 자동 안내 — 클릭 시 자동 입력'],
],col_widths=[4,12])

# ══════════════════ 9. 용어 정의 ══════════════════
h1('9. 용어 정의')
tbl(['용어','정의'],[
    ['Lot ID','공정 단위 식별자. 형식: 영숫자 6자리 + "." + 2자리 (예: ABC123.01)'],
    ['Net Die','웨이퍼 1장당 Good Die 수량 (과제 기준 고정값, 생산과제정보에서 등록)'],
    ['In(PCS)','투입 Wafer 수량'],
    ['In(EA)','투입 Die 수량 (= PCS × Net Die). ASSEMBLY·FT·DPS는 직접 입력'],
    ['Out(PCS)','출하 Wafer 수량 (Scrap 제외)'],
    ['Out(EA)','출하 Good Die 수량. ASSEMBLY·FT·DPS는 직접 입력'],
    ['Reject','In(EA) − Out(EA)'],
    ['YLD%','Out(EA) ÷ In(EA) × 100'],
    ['처리방향','다음step / 고객출하 / 보관 / scrap / 기타'],
    ['상태','진행중 / 완료 / Scrap'],
    ['SCM Flow','과제별 공정 순서. 생산과제정보 등록 화면에서 과제마다 자유롭게 구성'],
    ['Post-FAB','FAB 이후 공정 전체 (BUMP / CP / ASSEMBLY / FT / DPS)'],
],col_widths=[3.5,12.5])

out='/home/user/mj-automation-dashboard/docs/입출하관리_기능명세서_R5.docx'
doc.save(out)
print('saved:', out)
