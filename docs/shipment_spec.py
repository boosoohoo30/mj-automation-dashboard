from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 기본 스타일 ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_font(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=(0x1A,0x1D,0x23))
    # 하단 구분선
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '3B5BDB')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=11, color=(0x3B,0x5B,0xDB))

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10)

def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(level * 0.6)
    run = p.add_run(text)
    set_font(run, size=10)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run('※ ' + text)
    set_font(run, size=9, color=(0x5A,0x60,0x70))

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EEF2FF')
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, bold=True, size=9, color=(0x3B,0x5B,0xDB))

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9)

    # 컬럼 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('입출하관리 화면')
set_font(run, bold=True, size=22, color=(0x1A,0x1D,0x23))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('기능 명세서')
set_font(run, bold=True, size=16, color=(0x3B,0x5B,0xDB))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ERP 개발 요청용  |  2026-05')
set_font(run, size=10, color=(0x5A,0x60,0x70))

doc.add_page_break()


# ══════════════════════════════════════════════
# 1. 개요
# ══════════════════════════════════════════════
heading1('1. 개요')
body('반도체 제조 과정의 FAB → BUMP → CP → ASSEMBLY → FT 단계별 투입/출하 이력을 관리하는 화면입니다.')
body('과제(프로젝트) 단위로 관리되며, 각 공정 단계에서 IN/OUT 기록을 등록·조회할 수 있습니다.')

doc.add_paragraph()
add_table(
    ['항목', '내용'],
    [
        ['화면명', '입출하관리'],
        ['대상 공정', 'FAB / BUMP / CP / ASSEMBLY / FT'],
        ['관리 단위', 'MPW(EA), Wafer(PCS), Assembly·FT(EA)'],
        ['주요 기능', 'IN/OUT 등록, 잔여 추적, YLD 자동계산, Excel 업로드'],
    ],
    col_widths=[4, 12]
)


# ══════════════════════════════════════════════
# 2. 화면 구성
# ══════════════════════════════════════════════
heading1('2. 화면 구성')
body('화면은 좌측 과제 목록(사이드바)과 우측 상세 패널로 구성됩니다.')

doc.add_paragraph()
add_table(
    ['영역', '설명'],
    [
        ['과제 목록 (좌측)', '등록된 과제를 검색·선택. 과제명·AL Code 검색 가능'],
        ['과제 정보 바 (상단)', '선택된 과제의 AL Code·TM Code·Net Die·고객사 등 기본 정보 표시'],
        ['공정 탭', '해당 과제의 SCM Flow(예: FAB → CP)를 탭으로 표시. 탭 클릭 시 해당 단계 기록 조회'],
        ['기록 테이블 (우측 메인)', '선택 공정의 IN/OUT 기록 목록. IN 추가, OUT 등록, 편집, 삭제 가능'],
    ],
    col_widths=[4, 12]
)


# ══════════════════════════════════════════════
# 3. 과제(프로젝트) 분류
# ══════════════════════════════════════════════
heading1('3. 과제 분류')
body('과제는 아래 3가지 유형으로 구분되며, 유형별로 입출하 관리 방식이 다릅니다.')

doc.add_paragraph()
add_table(
    ['유형', 'SCM Flow', '관리 방식'],
    [
        ['MPW', 'FAB', 'EA 단위. IN = OUT 고정. Invoice(TSMC) 파싱 지원. 잔여/YLD 개념 없음'],
        ['Pilot / Production Wafer', 'FAB → CP, BUMP 등', 'Wafer 단장 단위 추적. Lot당 최대 25장. 잔여 장수 실시간 표시'],
        ['Assembly / FT 포함 과제', 'FAB → ASSEMBLY → FT 등', 'Assembly/FT 구간은 EA 직접 입력. Reject·YLD 자동계산'],
    ],
    col_widths=[3.5, 4.5, 8]
)


# ══════════════════════════════════════════════
# 4. FAB 단계 기능
# ══════════════════════════════════════════════
heading1('4. FAB 단계 기능')

heading2('4-1. FAB IN 등록')
add_table(
    ['구분', '입력 항목'],
    [
        ['MPW', 'Ver / 날짜 / Start Qty(EA) / Remark'],
        ['Wafer', 'Ver / 날짜 / Lot ID / Wafer 수량(최대 25) → Wafer No. 자동 할당(#1~N) / Remark'],
    ],
    col_widths=[3, 13]
)
note('Lot ID 형식: 영숫자 6자리 + "." + 2자리 (예: ABC123.01)')

heading2('4-2. FAB OUT 등록')
body('[MPW]')
bullet('TSMC Invoice 파싱 지원 — .txt 파일 업로드 또는 텍스트 붙여넣기')
bullet('파싱 결과: Invoice No. / 품목별 Chip ID·수량 자동 추출')
bullet('상태·처리방향 선택 후 저장')

doc.add_paragraph()
body('[Wafer]')
bullet('한 FAB IN에 대해 복수의 OUT 분배 등록 가능 (예: #1~3 고객출하 / #4~6 보관 / #7 scrap)')
bullet('Wafer Pool 시각화 — 이미 처리된 Wafer는 비활성 표시')
bullet('분배별 입력: Wafer No.(범위 표기 지원) / 처리방향 / 상태 / Reject / YLD% / Remark')
bullet('IN 행에 잔여 장수 배지 실시간 표시 (잔여 N장 / 완출하 / 미출하)')
note('Wafer No. 범위 입력: 쉼표·~ 사용 (예: 1~3,5,7~10). Excel 날짜 오인식 방지를 위해 ~ 권장')

heading2('4-3. 처리방향 옵션')
body('다음step / 고객출하 / 보관 / scrap / 기타')


# ══════════════════════════════════════════════
# 5. Post-FAB (CP / BUMP) 단계 기능
# ══════════════════════════════════════════════
heading1('5. Post-FAB 단계 기능 (CP / BUMP 등)')
body('Wafer 단위로 관리됩니다. Lot 단위로 등록하며 Wafer별 Good Die 수량을 입력합니다.')

heading2('5-1. Lot 등록 (IN)')
add_table(
    ['입력 항목', '설명'],
    [
        ['Lot ID', '직접 입력. 이전 단계 Lot ID 자동 가이드(클릭 시 자동 입력)'],
        ['Wafer No.', '범위 입력 (예: 1~3,5~6). FAB scrap으로 인한 비연속 번호 지원'],
        ['In(PCS)', 'Wafer 수량 자동계산'],
        ['In(EA)', 'Net Die × Wafer 수량 자동계산 (Net Die 등록 시)'],
        ['입고일 / 시작일 / 종료일', '날짜 입력'],
        ['상태 / 처리방향 / Remark', ''],
    ],
    col_widths=[4, 12]
)

heading2('5-2. Lot OUT (Wafer별 Good Die 입력)')
bullet('Wafer No. 입력 시 Wafer별 입력 그리드 자동 생성')
bullet('각 Wafer마다 Good Die 수량 및 Scrap 여부 입력')
bullet('Out(PCS) / Out(EA) / Reject / YLD% 자동계산 및 요약 표시')
bullet('Scrap 체크 시 해당 행 비활성화')


# ══════════════════════════════════════════════
# 6. Assembly / FT 단계 기능
# ══════════════════════════════════════════════
heading1('6. Assembly / FT 단계 기능')
body('다수 Wafer의 다이를 합산한 EA 단위로 관리합니다. Wafer 번호 개념 없음.')

add_table(
    ['입력 항목', '설명'],
    [
        ['Lot ID', '이전 단계 Lot ID 가이드 제공'],
        ['In(EA)', '직접 입력 (이전 단계 OUT EA 기준)'],
        ['Out(EA)', '직접 입력'],
        ['Reject', 'In(EA) − Out(EA) 자동계산'],
        ['YLD%', 'Out(EA) ÷ In(EA) × 100 자동계산'],
        ['입고일 / 시작일 / 종료일', '날짜 입력'],
        ['상태 / 처리방향 / Remark', ''],
    ],
    col_widths=[4, 12]
)


# ══════════════════════════════════════════════
# 7. Excel 업로드 기능
# ══════════════════════════════════════════════
heading1('7. Excel 업로드 기능')
body('공정 단계별 Excel 일괄 등록을 지원합니다. .xlsx / .xls / .csv 파일 업로드 가능.')

doc.add_paragraph()
add_table(
    ['모드', '헤더 컬럼'],
    [
        ['FAB IN',         'Ver | 날짜 | Lot ID | Wafer Qty (1-25) | Remark'],
        ['FAB OUT',        'Ver | 날짜 | Lot ID | Wafer No. | 처리방향 | 상태 | Remark'],
        ['Lot IN (CP등)',  'Lot ID | Wafer Qty | 입고일 | 시작일 | 종료일 | In(PCS) | In(EA) | 상태 | 처리방향 | Remark'],
        ['Lot OUT (CP등)', 'Lot ID | Wafer No. | Out Good Die (EA) | Is Scrap (Y/N) | Remark'],
        ['ASSY/FT IN',     'Lot ID | In(EA) | 입고일 | 시작일 | 종료일 | 상태 | 처리방향 | Remark'],
        ['ASSY/FT OUT',    'Lot ID | Out(EA) | 상태 | 처리방향 | Remark'],
    ],
    col_widths=[3.5, 12.5]
)
bullet('업로드 전 템플릿 미리보기 및 헤더 클립보드 복사 기능 제공')
bullet('파일 선택 후 데이터 미리보기 확인 → 저장 버튼으로 일괄 등록')
bullet('날짜 컬럼: Excel 날짜 시리얼 자동 변환 (YYYY-MM-DD)')
bullet('Wafer No. 범위: ~ 기호 사용 권장 (예: 1~3,5~20) — - 기호는 Excel이 날짜로 인식할 수 있음')


# ══════════════════════════════════════════════
# 8. 공통 기능
# ══════════════════════════════════════════════
heading1('8. 공통 기능')

add_table(
    ['기능', '설명'],
    [
        ['검색', '과제명·AL Code 키워드 검색'],
        ['편집', '등록된 IN/OUT/Lot 레코드 수정 (기존 값 자동 로드)'],
        ['삭제', '레코드 삭제 (확인 팝업 후 처리)'],
        ['Wafer별 상세 보기', 'CP/BUMP Lot 행에서 ▼ 버튼 클릭 시 Wafer별 Good Die·Scrap 현황 인라인 표시'],
        ['잔여 배지', 'FAB IN 행에 잔여 N장 / 완출하 / 미출하 상태 실시간 표시'],
        ['Lot ID 가이드', 'Lot 등록 시 이전 단계 Lot ID 자동 안내 — 클릭 시 자동 입력'],
    ],
    col_widths=[4, 12]
)


# ══════════════════════════════════════════════
# 9. 용어 정의
# ══════════════════════════════════════════════
heading1('9. 용어 정의')
add_table(
    ['용어', '정의'],
    [
        ['Lot ID',      '공정 단위 식별자. 형식: 영숫자 6자리 + "." + 2자리 (예: ABC123.01)'],
        ['Net Die',     '웨이퍼 1장당 Good Die 수량 (과제 기준 고정값)'],
        ['In(PCS)',     '투입 Wafer 수량'],
        ['In(EA)',      '투입 Die 수량 (= PCS × Net Die)'],
        ['Out(PCS)',    '출하 Wafer 수량 (scrap 제외)'],
        ['Out(EA)',     '출하 Good Die 수량'],
        ['Reject',      'In(EA) − Out(EA)'],
        ['YLD%',        'Out(EA) ÷ In(EA) × 100'],
        ['처리방향',    '다음step / 고객출하 / 보관 / scrap / 기타'],
        ['SCM Flow',    '과제별 공정 순서 (예: FAB → CP → FT)'],
    ],
    col_widths=[3.5, 12.5]
)


# 저장
out_path = '/home/user/mj-automation-dashboard/docs/입출하관리_기능명세서.docx'
doc.save(out_path)
print('saved:', out_path)
